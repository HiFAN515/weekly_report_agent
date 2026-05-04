"""
多格式导出器

功能：
  - Markdown 写文件
  - DOCX 导出（python-docx，含表格、样式）
  - HTML 导出（Markdown → HTML）
"""

from __future__ import annotations

import re
from datetime import date
from pathlib import Path


class Exporter:
    """周报导出器"""

    def __init__(self, output_dir: str | Path):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def export_markdown(self, content: str, week_start: date, week_end: date,
                        filename: str = None) -> Path:
        """导出为 Markdown 文件"""
        if not filename:
            filename = f"{week_start.isoformat()}_{week_end.isoformat()}_weekly_report.md"
        filepath = self.output_dir / filename
        filepath.write_text(content, encoding="utf-8")
        return filepath

    def export_docx(self, content: str, week_start: date, week_end: date,
                    filename: str = None) -> Path:
        """导出为 DOCX 文件（含表格、标题层级、列表）"""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not filename:
            filename = f"{week_start.isoformat()}_{week_end.isoformat()}_weekly_report.docx"
        filepath = self.output_dir / filename

        doc = Document()

        # 默认样式
        style = doc.styles['Normal']
        style.font.size = Pt(11)

        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].rstrip()

            # 标题
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)

            # 表格
            elif line.startswith("|") and "---" not in line:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    if "---" not in lines[i]:
                        table_lines.append(lines[i])
                    i += 1
                i -= 1  # 回退一行
                if table_lines:
                    self._add_table(doc, table_lines)

            # 任务列表
            elif line.startswith("- [ ] "):
                doc.add_paragraph(line[6:], style='List Bullet')
            elif line.startswith("- [x] "):
                p = doc.add_paragraph(line[6:], style='List Bullet')
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

            # 普通列表
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style='List Bullet')

            # 引用
            elif line.startswith("> "):
                p = doc.add_paragraph(line[2:])
                p.paragraph_format.left_indent = Pt(20)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # 分割线
            elif line.strip() == "---":
                doc.add_paragraph("─" * 50)

            # 粗体行
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line.strip("*"))
                run.bold = True

            # 普通文本
            elif line.strip():
                doc.add_paragraph(line)

            i += 1

        doc.save(str(filepath))
        return filepath

    def export_html(self, content: str, week_start: date, week_end: date,
                    filename: str = None) -> Path:
        """导出为 HTML 文件"""
        if not filename:
            filename = f"{week_start.isoformat()}_{week_end.isoformat()}_weekly_report.html"
        filepath = self.output_dir / filename

        html = self._markdown_to_html(content)
        filepath.write_text(html, encoding="utf-8")
        return filepath

    def _add_table(self, doc, table_lines: list[str]):
        """解析 Markdown 表格并添加到 docx"""
        rows = []
        for line in table_lines:
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells:
                rows.append(cells)

        if not rows:
            return

        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.style = 'Table Grid'

        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < len(table.rows[i].cells):
                    cell = table.rows[i].cells[j]
                    cell.text = cell_text
                    # 表头加粗
                    if i == 0:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.bold = True

    def _markdown_to_html(self, content: str) -> str:
        """Markdown 转 HTML"""
        import html as html_module

        lines = content.split("\n")
        body_parts = []
        in_list = False

        for line in lines:
            stripped = line.rstrip()

            # 标题
            if stripped.startswith("### "):
                if in_list:
                    body_parts.append("</ul>")
                    in_list = False
                body_parts.append(f"<h3>{html_module.escape(stripped[4:])}</h3>")
            elif stripped.startswith("## "):
                if in_list:
                    body_parts.append("</ul>")
                    in_list = False
                body_parts.append(f"<h2>{html_module.escape(stripped[3:])}</h2>")
            elif stripped.startswith("# "):
                if in_list:
                    body_parts.append("</ul>")
                    in_list = False
                body_parts.append(f"<h1>{html_module.escape(stripped[2:])}</h1>")

            # 表格
            elif stripped.startswith("|"):
                if in_list:
                    body_parts.append("</ul>")
                    in_list = False
                cells = [c.strip() for c in stripped.split("|")[1:-1]]
                if "---" not in stripped:
                    row = "".join(f"<td>{html_module.escape(c)}</td>" for c in cells)
                    body_parts.append(f"<tr>{row}</tr>")

            # 列表
            elif stripped.startswith("- "):
                if not in_list:
                    body_parts.append("<ul>")
                    in_list = True
                text = stripped[2:]
                text = re.sub(r'\[ \]', '☐', text)
                text = re.sub(r'\[x\]', '☑', text)
                body_parts.append(f"<li>{html_module.escape(text)}</li>")

            # 引用
            elif stripped.startswith("> "):
                if in_list:
                    body_parts.append("</ul>")
                    in_list = False
                body_parts.append(f"<blockquote>{html_module.escape(stripped[2:])}</blockquote>")

            # 分割线
            elif stripped == "---":
                if in_list:
                    body_parts.append("</ul>")
                    in_list = False
                body_parts.append("<hr>")

            # 空行
            elif not stripped:
                if in_list:
                    body_parts.append("</ul>")
                    in_list = False

            # 普通文本
            else:
                if in_list:
                    body_parts.append("</ul>")
                    in_list = False
                # 处理粗体
                text = html_module.escape(stripped)
                text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
                body_parts.append(f"<p>{text}</p>")

        if in_list:
            body_parts.append("</ul>")

        body = "\n".join(body_parts)

        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>周报</title>
<style>
  body {{ font-family: -apple-system, "Microsoft YaHei", sans-serif; max-width: 800px; margin: 40px auto; padding: 0 20px; line-height: 1.6; color: #333; }}
  h1 {{ border-bottom: 2px solid #eee; padding-bottom: 10px; }}
  h2 {{ color: #2c3e50; margin-top: 30px; }}
  h3 {{ color: #34495e; }}
  table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
  th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
  th {{ background: #f5f5f5; font-weight: bold; }}
  tr:nth-child(even) {{ background: #fafafa; }}
  blockquote {{ border-left: 4px solid #ddd; margin: 10px 0; padding: 5px 15px; color: #666; background: #fafafa; }}
  ul {{ padding-left: 20px; }}
  li {{ margin: 3px 0; }}
  hr {{ border: none; border-top: 1px solid #eee; margin: 20px 0; }}
  strong {{ color: #2c3e50; }}
</style>
</head>
<body>
{body}
</body>
</html>"""
